# -*- coding: utf-8 -*-
"""Генерация ТЗ на перенос сайта rm-info.ru -> rmgk.site (Tilda)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x85, 0x2A, 0x34)   # фирменный бордовый
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"

doc = Document()

# базовый стиль
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# поля страницы
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(2)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(14)
    return p


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def check(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("☐  ")
    r.font.color.rgb = ACCENT
    r.font.bold = True
    p.add_run(text)
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("▸ " + text)
    r.font.italic = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    return p


def kv_table(rows):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].width = Cm(6)
        cells[1].width = Cm(11)
        rk = cells[0].paragraphs[0].add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10.5)
        set_cell_bg(cells[0], "F4ECEC")
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ---------- ШАПКА ----------
title("Техническое задание")
subtitle("Перенос сайта rm-info.ru → rmgk.site (Tilda) с редактированием и ребрендингом")

# ---------- 1. ЦЕЛЬ ----------
h1("1. Цель")
para("Перенести действующий сайт с домена rm-info.ru на новый домен rmgk.site в рамках "
     "платформы Tilda, одновременно выполнив редактирование контента: полностью удалить "
     "упоминания бренда «Ресурс менеджмент» и убрать раздел мерча.")
note("Ключевое требование: перенос выполняется не «один в один», а с редактированием. "
     "Копирование без правок не принимается. Правки вносятся в процессе переноса, до публикации нового сайта.")

# ---------- 2. ИСХОДНЫЕ ДАННЫЕ ----------
h1("2. Исходные данные")
kv_table([
    ("Платформа", "Tilda (источник и приёмник — одна платформа)"),
    ("Домен-источник", "rm-info.ru"),
    ("Домен-приёмник", "rmgk.site"),
    ("Текущий статус источника", "Приостановлен Tilda — страницы доступны в редакторе личного кабинета"),
    ("Способ переноса", "Копирование проекта/страниц внутри Tilda + подключение нового домена"),
])

# ---------- 3. СОСТАВ РАБОТ ----------
h1("3. Состав работ")

para("3.1. Перенос структуры и дизайна", italic=False, color=ACCENT)
bullet("страницы сайта с сохранением вёрстки, блоков, стилей, шрифтов, анимаций и мобильной адаптации.", lead="Перенести все ")
bullet("Сохранить логику навигации: меню, футер, внутренние ссылки, «хлебные крошки», кнопки.")
bullet("Перенести вложенные элементы: формы, всплывающие окна, zero-блоки, встроенный код, HTML-виджеты.")

para("3.2. Ребрендинг — удаление бренда «Ресурс менеджмент» (сплошная зачистка)", color=ACCENT)
para("Удалить/заменить любые упоминания во всех формах написания: «Ресурс менеджмент», "
     "«Ресурс-менеджмент», «Resource Management» и производные. Проверить и вычистить:")
bullet("тексты, заголовки, подзаголовки, кнопки;")
bullet("логотип и все графические носители бренда (favicon, водяные знаки на изображениях, баннеры);")
bullet("alt изображений, title, meta-теги, Open Graph, микроразметку;")
bullet("футер, копирайт, юридические реквизиты, название юрлица — заменить на реквизиты ООО «РМГК» (см. раздел 4);")
bullet("e-mail, телефоны, ссылки на соцсети, содержащие старый бренд;")
bullet("адреса страниц (URL/slug), названия страниц в личном кабинете Tilda;")
bullet("тексты форм, авто-ответов и уведомлений о заявках.")
note("Исключение: бренд «РМ Групп» сохраняется — его удалять и заменять не нужно. "
     "Зачистке подлежит только «Ресурс менеджмент» и его варианты написания.")

para("3.3. Мерч — не переносить", color=ACCENT)
bullet("Раздел/страницу мерча на новый сайт не переносить вообще: не копировать каталог, карточки товаров, цены, кнопки заказа, формы и изображения продукции.")
bullet("Не переносить пункт(ы) меню и любые внутренние ссылки, ведущие на мерч.")
bullet("Убедиться, что на новом сайте не осталось «висящих» ссылок, кнопок и пустых блоков, связанных с мерчем.")
note("Мерч исключается из переноса полностью. На новом сайте rmgk.site упоминаний и элементов мерча быть не должно.")

para("3.4. Домен, техника, SEO", color=ACCENT)
bullet("Подключить домен rmgk.site к новому проекту Tilda, настроить DNS (A-запись / CNAME по инструкции Tilda), выпустить SSL-сертификат (HTTPS).")
bullet("Переопределить title / description / OG на каждой странице без старого бренда.")
bullet("Заменить favicon на новый.")
bullet("Пересобрать и опубликовать sitemap.xml, проверить robots.txt (снять возможные noindex).")
bullet("Перенастроить приёмники форм (почта / CRM / Telegram / Google-таблица) на актуальные реквизиты нового проекта.")
bullet("По возможности настроить 301-редиректы со старых URL (если старый домен будет доступен).")

# ---------- 4. РЕКВИЗИТЫ / ДОСТУПЫ ----------
h1("4. Реквизиты нового юрлица и доступы")
para("Все брендовые элементы старого сайта заменяются на реквизиты нового юридического лица:")
kv_table([
    ("Наименование", "ООО «РМГК»"),
    ("ИНН", "9701311770"),
    ("КПП", "770101001"),
    ("Юридический адрес", "г. Москва, Большая Почтовая ул., д. 26в стр. 1"),
    ("E-mail", "info@rmgk.ru"),
    ("Копирайт в футере", "© ООО «РМГК». Все права защищены."),
])
para("Заказчик дополнительно предоставляет:")
bullet("Доступ к личному кабинету Tilda (источник и приёмник) либо тариф, позволяющий перенос.")
bullet("Доступ к управлению доменом rmgk.site (панель регистратора / DNS).")
bullet("Логотип и favicon нового бренда (или указание на их отсутствие/заглушку).")

# ---------- 5. ПРИЁМКА ----------
h1("5. Приёмка (чек-лист)")
para("Работа принимается, если:")
check("все страницы открываются на rmgk.site по HTTPS без ошибок;")
check("поиск по сайту и по коду страниц не находит ни одного упоминания старого бренда (тексты, alt, meta, URL, футер);")
check("раздел мерча и все ссылки на него отсутствуют в навигации и теле страниц;")
check("формы отправляют заявки на указанные заказчиком приёмники (проверено тестовой заявкой);")
check("мобильная версия корректна, битых ссылок и пустых блоков нет;")
check("favicon, title и description обновлены на всех страницах.")

# ---------- 6. СРОКИ ----------
h1("6. Сроки и формат сдачи")
bullet("5 (пять) рабочих дней с момента предоставления доступов.", lead="Срок: ")
bullet("Сдача: ссылка на опубликованный сайт rmgk.site + краткий отчёт о внесённых правках "
       "(список удалённых брендовых элементов и скрытых страниц мерча).")

out = "/Users/daria/Desktop/Сlaude Code/ТЗ_перенос_rmgk.site.docx"
doc.save(out)
print("Saved:", out)
